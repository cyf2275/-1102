from __future__ import annotations

import csv
import json
import math
import statistics
from pathlib import Path
from typing import Any

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path.cwd() / "cloud_results" / "A_20260619_selfbuilt_dataset_paper_experiments"
ASSET_DIR = ROOT / "paper_summary_assets"
REPORT_MD = ROOT / "自建数据集单帧相位后验扩散_RCPC_实验总结.md"
REPORT_DOCX = ROOT / "自建数据集单帧相位后验扩散_RCPC_实验总结.docx"
SUMMARY_JSON = ROOT / "paper_summary_metrics.json"
SUMMARY_CSV = ROOT / "paper_summary_metrics.csv"
SIGNIFICANCE_JSON = ROOT / "paper_per_sample_significance.json"
SIGNIFICANCE_CSV = ROOT / "paper_per_sample_significance.csv"
PER_SAMPLE_MEAN_CSV = ROOT / "paper_per_sample_mean_rmse.csv"


FORMAL_DIRECT = ROOT / "A_20260619_formal_strong_backbone_direct_seed012"
FORMAL_SELECTOR = ROOT / "A_20260619_formal_attention_unet_ours_selector_seed012"
QUICK_BASELINE = ROOT / "A_20260618_single_frame3d_baseline_comparison_quick1seed"
PAPER_READY = ROOT / "A_20260618_paper_ready_anchor_ablation"
DIRECT_PLY = ROOT / "A_20260619_direct_baseline_obj061_pose02_ply"

CODE_PATHS = {
    "正式 strong backbone 训练与评估": Path.cwd()
    / "diffusion_fpp_v5"
    / "train_single_frame3d_backbone_baselines.py",
    "Attention U-Net + ours selector 训练/评估": Path.cwd()
    / "diffusion_fpp_v5"
    / "train_refined_xphase_reliability_selector.py",
    "x 相位后验扩散与 RCPC 主流程": Path.cwd()
    / "diffusion_fpp_v5"
    / "train_single_frame3d_xphase_diffusion_rcpc.py",
    "最佳方法二维重建可视化": Path.cwd()
    / "diffusion_fpp_v5"
    / "make_best_anchor_reconstruction_visuals.py",
    "ours/anchor 严格相机反投影 PLY 导出": Path.cwd()
    / "diffusion_fpp_v5"
    / "export_best_anchor_pose_ply.py",
    "direct baseline 严格相机反投影 PLY 导出": Path.cwd()
    / "diffusion_fpp_v5"
    / "export_single_frame3d_direct_baseline_ply.py",
    "样本级显著性分析": Path.cwd()
    / "tools"
    / "add_selfbuilt_dataset_supplementary_analysis_20260619.py",
    "本文档、图表和 Word 生成": Path.cwd()
    / "tools"
    / "build_selfbuilt_dataset_report_20260619.py",
}


def f4(x: Any) -> str:
    if x is None:
        return "-"
    if isinstance(x, float) and math.isnan(x):
        return "-"
    if isinstance(x, (int, float)):
        return f"{x:.4f}"
    return str(x)


def pct(a: float, b: float) -> float:
    return (a - b) / a * 100.0


def read_csv_rows(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def read_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8"))


def mean_std(values: list[float]) -> dict[str, float]:
    return {
        "mean": statistics.mean(values),
        "std": statistics.stdev(values) if len(values) > 1 else 0.0,
        "n": len(values),
    }


def aggregate_direct() -> dict[str, dict[str, dict[str, float]]]:
    rows = read_csv_rows(FORMAL_DIRECT / "baseline_comparison_quick1seed_summary.csv")
    grouped: dict[str, dict[str, list[float]]] = {}
    for row in rows:
        method = row["method"]
        grouped.setdefault(method, {"test": [], "ood": [], "val": []})
        grouped[method]["test"].append(float(row["test_object_rmse"]))
        grouped[method]["ood"].append(float(row["ood_object_rmse"]))
        grouped[method]["val"].append(float(row["best_val_object_rmse"]))
    return {
        method: {split: mean_std(values) for split, values in split_values.items()}
        for method, split_values in grouped.items()
    }


def aggregate_quick() -> list[dict[str, Any]]:
    path = QUICK_BASELINE / "baseline_comparison_quick1seed_summary.csv"
    rows = read_csv_rows(path)
    out = []
    for row in rows:
        if row["kind"] != "quick_1seed_baseline":
            continue
        out.append(
            {
                "method": row["method"],
                "test": float(row["test_object_rmse"]),
                "ood": float(row["ood_object_rmse"]),
                "val": float(row["best_val_object_rmse"]),
                "epochs": int(float(row["epochs"])),
            }
        )
    return sorted(out, key=lambda r: r["test"])


def selector_aggregate() -> dict[str, Any]:
    data = read_json(FORMAL_SELECTOR / "formal_attention_unet_ours_selector_summary.json")
    return data["aggregate"]


def significance_rows() -> list[dict[str, Any]]:
    if not SIGNIFICANCE_CSV.exists():
        return []
    rows: list[dict[str, Any]] = []
    for row in read_csv_rows(SIGNIFICANCE_CSV):
        row = dict(row)
        for key in [
            "n",
            "baseline_mean",
            "candidate_mean",
            "mean_improvement_mm",
            "median_improvement_mm",
            "relative_improvement_percent",
            "wins",
            "losses",
            "ties",
            "win_rate",
            "sign_test_p_two_sided",
        ]:
            if key in row:
                try:
                    row[key] = float(row[key])
                except ValueError:
                    pass
        if "n" in row:
            row["n"] = int(row["n"])
        for key in ["wins", "losses", "ties"]:
            if key in row:
                row[key] = int(row[key])
        rows.append(row)
    return rows


def write_metric_files(
    direct: dict[str, Any], selector: dict[str, Any], quick: list[dict[str, Any]]
) -> None:
    payload = {
        "local_root": str(ROOT),
        "formal_direct": direct,
        "formal_attention_unet_ours_selector": selector,
        "quick_1seed_screening": quick,
        "notes": [
            "Formal direct uses 80 epochs with best-validation checkpoint selection.",
            "Quick screening uses one seed and 40 epochs; it is not the final paper table.",
            "Primary metric is object-mask RMSE in mm; valid-mask RMSE is auxiliary.",
        ],
    }
    SUMMARY_JSON.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

    rows = []
    for method, split_data in direct.items():
        rows.append(
            {
                "group": "formal_direct",
                "method": method,
                "test_mean": split_data["test"]["mean"],
                "test_std": split_data["test"]["std"],
                "ood_mean": split_data["ood"]["mean"],
                "ood_std": split_data["ood"]["std"],
            }
        )
    for method in ["base", "x_phase", "anchor", "refined", "sample_rcpc", "rule", "mlp", "true_x_oracle"]:
        rows.append(
            {
                "group": "formal_attention_unet_ours_selector",
                "method": method,
                "test_mean": selector["test"][method]["mean"],
                "test_std": selector["test"][method]["std"],
                "ood_mean": selector["ood"][method]["mean"],
                "ood_std": selector["ood"][method]["std"],
            }
        )
    with SUMMARY_CSV.open("w", encoding="utf-8", newline="") as f:
        writer = csv.DictWriter(
            f,
            fieldnames=["group", "method", "test_mean", "test_std", "ood_mean", "ood_std"],
        )
        writer.writeheader()
        writer.writerows(rows)


def set_plot_style() -> None:
    plt.rcParams.update(
        {
            "font.family": "DejaVu Sans",
            "axes.spines.top": False,
            "axes.spines.right": False,
            "axes.grid": True,
            "grid.alpha": 0.25,
            "figure.dpi": 160,
            "savefig.dpi": 220,
        }
    )


def plot_formal_comparison(direct: dict[str, Any], selector: dict[str, Any]) -> Path:
    labels = [
        "Attention U-Net\nDirect",
        "UNet++\nDirect",
        "Base+x\nAnchor",
        "Rule\nFinal",
        "MLP\nFinal",
    ]
    test_means = [
        direct["attention_unet"]["test"]["mean"],
        direct["unetpp"]["test"]["mean"],
        selector["test"]["anchor"]["mean"],
        selector["test"]["rule"]["mean"],
        selector["test"]["mlp"]["mean"],
    ]
    test_stds = [
        direct["attention_unet"]["test"]["std"],
        direct["unetpp"]["test"]["std"],
        selector["test"]["anchor"]["std"],
        selector["test"]["rule"]["std"],
        selector["test"]["mlp"]["std"],
    ]
    ood_means = [
        direct["attention_unet"]["ood"]["mean"],
        direct["unetpp"]["ood"]["mean"],
        selector["ood"]["anchor"]["mean"],
        selector["ood"]["rule"]["mean"],
        selector["ood"]["mlp"]["mean"],
    ]
    ood_stds = [
        direct["attention_unet"]["ood"]["std"],
        direct["unetpp"]["ood"]["std"],
        selector["ood"]["anchor"]["std"],
        selector["ood"]["rule"]["std"],
        selector["ood"]["mlp"]["std"],
    ]
    x = np.arange(len(labels))
    w = 0.36
    fig, ax = plt.subplots(figsize=(8.8, 4.8))
    ax.bar(x - w / 2, test_means, w, yerr=test_stds, capsize=4, label="Test", color="#4C78A8")
    ax.bar(x + w / 2, ood_means, w, yerr=ood_stds, capsize=4, label="OOD 61-64", color="#F58518")
    ax.set_ylabel("Object-mask RMSE (mm)")
    ax.set_title("Formal 3-seed comparison, 80 epochs / best-val checkpoint")
    ax.set_xticks(x, labels)
    ax.legend(frameon=False)
    ax.set_ylim(1.25, 2.15)
    for i, v in enumerate(test_means):
        ax.text(i - w / 2, v + 0.03, f"{v:.3f}", ha="center", va="bottom", fontsize=8)
    for i, v in enumerate(ood_means):
        ax.text(i + w / 2, v + 0.03, f"{v:.3f}", ha="center", va="bottom", fontsize=8)
    fig.tight_layout()
    out = ASSET_DIR / "formal_test_ood_comparison.png"
    fig.savefig(out)
    plt.close(fig)
    return out


def plot_selector_ablation(selector: dict[str, Any]) -> Path:
    methods = ["base", "x_phase", "anchor", "refined", "sample_rcpc", "rule", "mlp", "true_x_oracle"]
    labels = ["Direct\nbase", "X phase", "Base+x\nanchor", "Diffusion\ncandidate", "Sample\nRCPC", "Rule\nfinal", "MLP\nfinal", "True x\noracle"]
    x = np.arange(len(methods))
    w = 0.36
    test = [selector["test"][m]["mean"] for m in methods]
    test_std = [selector["test"][m]["std"] for m in methods]
    ood = [selector["ood"][m]["mean"] for m in methods]
    ood_std = [selector["ood"][m]["std"] for m in methods]
    fig, ax = plt.subplots(figsize=(10.5, 4.8))
    ax.bar(x - w / 2, test, w, yerr=test_std, capsize=3, label="Test", color="#54A24B")
    ax.bar(x + w / 2, ood, w, yerr=ood_std, capsize=3, label="OOD 61-64", color="#E45756")
    ax.set_ylabel("Object-mask RMSE (mm)")
    ax.set_title("Ablation of Attention U-Net base + phase posterior evidence")
    ax.set_xticks(x, labels)
    ax.legend(frameon=False)
    ax.set_ylim(1.25, 2.05)
    fig.tight_layout()
    out = ASSET_DIR / "selector_ablation_test_ood.png"
    fig.savefig(out)
    plt.close(fig)
    return out


def plot_quick_screening(quick: list[dict[str, Any]]) -> Path:
    labels = [r["method"].replace("_", " ") for r in quick]
    test = [r["test"] for r in quick]
    ood = [r["ood"] for r in quick]
    x = np.arange(len(labels))
    w = 0.36
    fig, ax = plt.subplots(figsize=(9.2, 4.8))
    ax.bar(x - w / 2, test, w, label="Test", color="#72B7B2")
    ax.bar(x + w / 2, ood, w, label="OOD 61-64", color="#B279A2")
    ax.set_ylabel("Object-mask RMSE (mm)")
    ax.set_title("Quick 1-seed baseline screening, 40 epochs")
    ax.set_xticks(x, labels, rotation=18, ha="right")
    ax.legend(frameon=False)
    fig.tight_layout()
    out = ASSET_DIR / "quick_baseline_screening.png"
    fig.savefig(out)
    plt.close(fig)
    return out


def text_font(size: int = 22) -> ImageFont.ImageFont:
    candidates = [
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibri.ttf"),
        Path("C:/Windows/Fonts/msyh.ttc"),
    ]
    for c in candidates:
        if c.exists():
            return ImageFont.truetype(str(c), size=size)
    return ImageFont.load_default()


def make_contact_sheet(items: list[tuple[str, Path]], out: Path, thumb_width: int = 760) -> Path:
    font = text_font(24)
    title_h = 42
    margin = 18
    thumbs: list[tuple[str, Image.Image]] = []
    for label, path in items:
        img = Image.open(path).convert("RGB")
        scale = thumb_width / img.width
        img = img.resize((thumb_width, max(1, int(img.height * scale))), Image.LANCZOS)
        thumbs.append((label, img))
    cols = 2
    rows = math.ceil(len(thumbs) / cols)
    cell_h = max(img.height for _, img in thumbs) + title_h + margin
    w = cols * thumb_width + (cols + 1) * margin
    h = rows * cell_h + margin
    canvas = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(canvas)
    for idx, (label, img) in enumerate(thumbs):
        r, c = divmod(idx, cols)
        x = margin + c * (thumb_width + margin)
        y = margin + r * cell_h
        draw.text((x, y), label, fill=(20, 20, 20), font=font)
        canvas.paste(img, (x, y + title_h))
    canvas.save(out, quality=95)
    return out


def make_visual_contact_sheet() -> Path:
    items = [
        (
            "Attention U-Net direct, 80ep seed0, obj061 pose02",
            FORMAL_DIRECT
            / "attention_unet_seed0"
            / "visualizations"
            / "ood"
            / "ood_01_new0612_obj061_pose02.png",
        ),
        (
            "UNet++ direct, 80ep seed0, obj061 pose02",
            FORMAL_DIRECT
            / "unetpp_seed0"
            / "visualizations"
            / "ood"
            / "ood_01_new0612_obj061_pose02.png",
        ),
        (
            "Ours detailed fullchain, seed0, obj061 pose02",
            PAPER_READY
            / "best_method_visuals_fullchain_base_x_mean_seed0"
            / "ood"
            / "ood_03_new0612_obj061_pose02.png",
        ),
        (
            "Ours overview, selected test/OOD samples",
            PAPER_READY
            / "best_method_visuals_fullchain_base_x_mean_seed0"
            / "best_method_reconstruction_overview.png",
        ),
    ]
    return make_contact_sheet(items, ASSET_DIR / "visual_comparison_contact_sheet.png")


def read_ply_points(path: Path, max_points: int = 9000) -> np.ndarray:
    with path.open("r", encoding="utf-8", errors="replace") as f:
        header = []
        n_vertices = None
        for line in f:
            header.append(line.rstrip("\n"))
            if line.startswith("element vertex"):
                n_vertices = int(line.split()[-1])
            if line.strip() == "end_header":
                break
        if n_vertices is None:
            raise ValueError(f"No vertex count in {path}")
        pts = []
        stride = max(1, n_vertices // max_points)
        for i, line in enumerate(f):
            if i % stride != 0:
                continue
            parts = line.split()
            if len(parts) >= 3:
                pts.append((float(parts[0]), float(parts[1]), float(parts[2])))
            if len(pts) >= max_points:
                break
    return np.asarray(pts, dtype=np.float32)


def set_axes_equal(ax: Any, pts_list: list[np.ndarray]) -> None:
    pts = np.concatenate(pts_list, axis=0)
    mins = pts.min(axis=0)
    maxs = pts.max(axis=0)
    centers = (mins + maxs) / 2.0
    radius = float((maxs - mins).max() / 2.0)
    ax.set_xlim(centers[0] - radius, centers[0] + radius)
    ax.set_ylim(centers[1] - radius, centers[1] + radius)
    ax.set_zlim(centers[2] - radius, centers[2] + radius)


def make_pointcloud_preview(kind: str, files: list[tuple[str, Path]], out: Path, cols: int = 2) -> Path:
    pts_list = [read_ply_points(path) for _, path in files]
    rows = int(math.ceil(len(files) / cols))
    fig = plt.figure(figsize=(5.1 * cols, 4.2 * rows))
    for idx, ((label, _), pts) in enumerate(zip(files, pts_list), start=1):
        ax = fig.add_subplot(rows, cols, idx, projection="3d")
        z = pts[:, 2]
        ax.scatter(pts[:, 0], pts[:, 1], pts[:, 2], c=z, s=1, cmap="viridis", alpha=0.85)
        ax.set_title(label, fontsize=10)
        ax.set_xlabel("X")
        ax.set_ylabel("Y")
        ax.set_zlabel("Z")
        ax.view_init(elev=24, azim=-62)
        set_axes_equal(ax, pts_list)
    fig.suptitle(kind, fontsize=13)
    fig.tight_layout()
    fig.savefig(out)
    plt.close(fig)
    return out


def make_pointcloud_previews() -> dict[str, Path]:
    cam_dir = PAPER_READY / "pointcloud_obj061_pose02_camera_projective"
    cam_files = [
        ("GT strict camera", cam_dir / "new0612_obj061_pose02_gt_camera_projective.ply"),
        ("Anchor strict camera", cam_dir / "new0612_obj061_pose02_anchor_base_x_mean_camera_projective.ply"),
        ("Rule final strict camera", cam_dir / "new0612_obj061_pose02_ours_rule_final_camera_projective.ply"),
        ("MLP final strict camera", cam_dir / "new0612_obj061_pose02_ours_mlp_final_camera_projective.ply"),
    ]
    strict = make_pointcloud_preview(
        "obj061 pose02 PLY preview, strict camera projective backprojection",
        cam_files,
        ASSET_DIR / "pointcloud_obj061_pose02_camera_projective_preview.png",
    )

    vis_dir = PAPER_READY / "pointcloud_obj061_pose02"
    vis_files = [
        ("GT visual z10", vis_dir / "new0612_obj061_pose02_gt_z10_visual.ply"),
        ("Anchor visual z10", vis_dir / "new0612_obj061_pose02_anchor_base_x_mean_z10_visual.ply"),
        ("Rule final visual z10", vis_dir / "new0612_obj061_pose02_ours_rule_final_z10_visual.ply"),
        ("MLP final visual z10", vis_dir / "new0612_obj061_pose02_ours_mlp_final_z10_visual.ply"),
    ]
    visual = make_pointcloud_preview(
        "obj061 pose02 PLY preview, visual z10 shape inspection only",
        vis_files,
        ASSET_DIR / "pointcloud_obj061_pose02_visual_z10_preview.png",
    )
    out = {"ply_strict": strict, "ply_visual": visual}

    direct_cam_files = [
        ("GT strict camera", cam_dir / "new0612_obj061_pose02_gt_camera_projective.ply"),
        (
            "Attention direct",
            DIRECT_PLY
            / "attention_unet_seed0"
            / "new0612_obj061_pose02_attention_unet_direct_seed0_camera_projective.ply",
        ),
        (
            "UNet++ direct",
            DIRECT_PLY / "unetpp_seed0" / "new0612_obj061_pose02_unetpp_direct_seed0_camera_projective.ply",
        ),
        ("Base+x anchor", cam_dir / "new0612_obj061_pose02_anchor_base_x_mean_camera_projective.ply"),
        ("Rule final", cam_dir / "new0612_obj061_pose02_ours_rule_final_camera_projective.ply"),
        ("MLP final", cam_dir / "new0612_obj061_pose02_ours_mlp_final_camera_projective.ply"),
    ]
    if all(path.exists() for _, path in direct_cam_files):
        out["ply_direct_strict"] = make_pointcloud_preview(
            "obj061 pose02 strict camera PLY: direct baselines vs ours",
            direct_cam_files,
            ASSET_DIR / "pointcloud_obj061_pose02_direct_vs_ours_camera_projective.png",
            cols=3,
        )

    direct_vis_files = [
        ("GT visual z10", vis_dir / "new0612_obj061_pose02_gt_z10_visual.ply"),
        (
            "Attention direct",
            DIRECT_PLY / "attention_unet_seed0" / "new0612_obj061_pose02_attention_unet_direct_seed0_z10_visual.ply",
        ),
        ("UNet++ direct", DIRECT_PLY / "unetpp_seed0" / "new0612_obj061_pose02_unetpp_direct_seed0_z10_visual.ply"),
        ("Base+x anchor", vis_dir / "new0612_obj061_pose02_anchor_base_x_mean_z10_visual.ply"),
        ("Rule final", vis_dir / "new0612_obj061_pose02_ours_rule_final_z10_visual.ply"),
        ("MLP final", vis_dir / "new0612_obj061_pose02_ours_mlp_final_z10_visual.ply"),
    ]
    if all(path.exists() for _, path in direct_vis_files):
        out["ply_direct_visual"] = make_pointcloud_preview(
            "obj061 pose02 visual z10 PLY: direct baselines vs ours",
            direct_vis_files,
            ASSET_DIR / "pointcloud_obj061_pose02_direct_vs_ours_visual_z10.png",
            cols=3,
        )
    return out


def md_table(headers: list[str], rows: list[list[Any]]) -> str:
    out = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    for row in rows:
        out.append("| " + " | ".join(str(x) for x in row) + " |")
    return "\n".join(out)


def build_markdown(direct: dict[str, Any], selector: dict[str, Any], quick: list[dict[str, Any]], assets: dict[str, Path]) -> str:
    att = direct["attention_unet"]
    unetpp = direct["unetpp"]
    test_gain = pct(att["test"]["mean"], selector["test"]["mlp"]["mean"])
    ood_gain_rule = pct(att["ood"]["mean"], selector["ood"]["rule"]["mean"])
    ood_gain_unetpp = pct(unetpp["ood"]["mean"], selector["ood"]["rule"]["mean"])
    sig_rows_all = significance_rows()
    sig_pick = []
    for split, baseline, candidate in [
        ("test", "attention_unet", "mlp"),
        ("test", "unetpp", "mlp"),
        ("ood", "attention_unet", "rule"),
        ("ood", "unetpp", "rule"),
    ]:
        match = next(
            (
                r
                for r in sig_rows_all
                if r.get("split") == split and r.get("baseline") == baseline and r.get("candidate") == candidate
            ),
            None,
        )
        if match is not None:
            sig_pick.append(match)

    quick_rows = [
        [r["method"], r["epochs"], f4(r["val"]), f4(r["test"]), f4(r["ood"])]
        for r in quick
    ]
    formal_rows = [
        [
            "Attention U-Net direct",
            f"{f4(att['test']['mean'])} ± {f4(att['test']['std'])}",
            f"{f4(att['ood']['mean'])} ± {f4(att['ood']['std'])}",
            "强 backbone 直接深度回归",
        ],
        [
            "UNet++ direct",
            f"{f4(unetpp['test']['mean'])} ± {f4(unetpp['test']['std'])}",
            f"{f4(unetpp['ood']['mean'])} ± {f4(unetpp['ood']['std'])}",
            "强 backbone 直接深度回归",
        ],
        [
            "Ours: base+x anchor",
            f"{f4(selector['test']['anchor']['mean'])} ± {f4(selector['test']['anchor']['std'])}",
            f"{f4(selector['ood']['anchor']['mean'])} ± {f4(selector['ood']['anchor']['std'])}",
            "基础锚点，不含最终选择器",
        ],
        [
            "Ours: Rule final",
            f"{f4(selector['test']['rule']['mean'])} ± {f4(selector['test']['rule']['std'])}",
            f"{f4(selector['ood']['rule']['mean'])} ± {f4(selector['ood']['rule']['std'])}",
            "OOD 最稳",
        ],
        [
            "Ours: MLP final",
            f"{f4(selector['test']['mlp']['mean'])} ± {f4(selector['test']['mlp']['std'])}",
            f"{f4(selector['ood']['mlp']['mean'])} ± {f4(selector['ood']['mlp']['std'])}",
            "普通 test 最好",
        ],
    ]
    ablation_rows = []
    name_map = {
        "base": "Direct base",
        "x_phase": "X phase candidate",
        "anchor": "Base+x mean anchor",
        "refined": "Phase posterior diffusion candidate",
        "sample_rcpc": "Sample RCPC",
        "rule": "Rule final",
        "mlp": "MLP final",
        "true_x_oracle": "True x phase oracle",
    }
    for key in ["base", "x_phase", "anchor", "refined", "sample_rcpc", "rule", "mlp", "true_x_oracle"]:
        ablation_rows.append(
            [
                name_map[key],
                f"{f4(selector['test'][key]['mean'])} ± {f4(selector['test'][key]['std'])}",
                f"{f4(selector['ood'][key]['mean'])} ± {f4(selector['ood'][key]['std'])}",
            ]
        )

    significance_table = []
    for row in sig_pick:
        significance_table.append(
            [
                row["split"],
                f"{row['baseline']} -> {row['candidate']}",
                f"{row['mean_improvement_mm']:.4f}",
                f"{row['relative_improvement_percent']:.1f}%",
                f"{row['wins']}/{row['n']}",
                f"{row['sign_test_p_two_sided']:.4g}",
            ]
        )

    rel = lambda p: p.relative_to(ROOT).as_posix()
    significance_section = ""
    if significance_table:
        significance_section = f"""
## 样本级显著性分析

统计方式：先对每个样本的 3 个 seed RMSE 取平均，再做 paired comparison。`wins/n` 表示候选方法在多少个样本上低于 baseline。

{md_table(["split", "paired comparison", "mean improvement mm", "relative", "wins/n", "sign-test p"], significance_table)}

![Per-sample paired improvement]({rel(assets['significance'])})

补充文件：

- `{SIGNIFICANCE_CSV}`
- `{SIGNIFICANCE_JSON}`
- `{PER_SAMPLE_MEAN_CSV}`
"""

    direct_ply_section = ""
    if "ply_direct_strict" in assets:
        direct_ply_section = f"""
新增 direct baseline 三维点云对比：

![Direct vs ours strict PLY]({rel(assets['ply_direct_strict'])})

![Direct vs ours visual PLY]({rel(assets['ply_direct_visual'])})

direct baseline PLY 文件位于：

`{DIRECT_PLY}`
"""

    code_rows = [[name, str(path)] for name, path in CODE_PATHS.items()]

    md = f"""# 自建数据集单帧相位后验扩散与 RCPC 实验总结

生成时间：2026-06-19

本报告整理服务器自建真实数据集实验结果，重点回答三个问题：第一，强 backbone 直接深度回归是否已经足够；第二，单帧 x 相位证据、phase posterior diffusion 和 RCPC/selector 是否仍有价值；第三，当前 80 epoch、3 seeds 的结果是否够作为论文实验表。

## 结论摘要

- 正式 3-seed 结果显示，Attention U-Net direct 的普通 test object RMSE 为 `{f4(att['test']['mean'])} ± {f4(att['test']['std'])}` mm，UNet++ direct 为 `{f4(unetpp['test']['mean'])} ± {f4(unetpp['test']['std'])}` mm。
- 接入我们的方法后，MLP final 在普通 test 达到 `{f4(selector['test']['mlp']['mean'])} ± {f4(selector['test']['mlp']['std'])}` mm，相对 Attention U-Net direct 改善约 `{test_gain:.1f}%`。
- 在异材质 61-64 OOD 上，Rule final 达到 `{f4(selector['ood']['rule']['mean'])} ± {f4(selector['ood']['rule']['std'])}` mm，相对 Attention U-Net direct 改善约 `{ood_gain_rule:.1f}%`，相对 UNet++ direct 改善约 `{ood_gain_unetpp:.1f}%`。
- phase posterior diffusion 的 refined depth candidate 单独不是最优，说明扩散不能简单写成“直接修深度一定更好”；更合理的论文表述是：扩散作为 x 相位后验候选，结合 anchor 与 selector/RCPC 后产生稳定增益。
- 80 epoch 对本轮正式对比是够用的：每个正式方法均为 seeds 0/1/2，训练到 80 epoch，并使用 best-val checkpoint；40 epoch 只作为快速筛选，不应写进最终主表。

## 数据与训练口径

- 数据：服务器 `/root/autodl-tmp/single_frame_3d_dataset_v1_upload_smalltest`，目标为 `depth_z`。
- 划分：train 352，val 80，test 31；额外 OOD 为 61-64 异材质样本，共 12。
- 合法 test-time 输入：`input_vertical_0120.bmp` 的单帧条纹图，以及从该单帧派生的物理/相位证据；真实 `phase_x/phase_y` 只能用于监督、诊断或 oracle，不作为 test-time 输入。
- 主指标：object-mask RMSE，单位 mm；valid-mask RMSE 作为辅助。
- 正式 direct backbone：Attention U-Net 和 UNet++，seeds 0/1/2，80 epoch，480x640，batch size 2，gradient accumulation 2，eval every 5 epoch，保存 best-val checkpoint。
- 我们的方法：使用正式 Attention U-Net direct checkpoint 作为 base，接入 x phase evidence、phase posterior diffusion/refined candidate、base+x mean anchor，并用 Rule 或 MLP selector 得到 final depth。

## 使用的方法

### 快速筛选

先用 1 seed、40 epoch 对多种单帧直接深度回归 baseline 做筛选，包括 U-Net、ResUNet、Attention U-Net、UNet++、Pix2Pix、MPS-XNet-style。这个阶段的意义是判断强 backbone 候选，不作为论文最终表。

{md_table(["方法", "epoch", "val RMSE", "test RMSE", "OOD RMSE"], quick_rows)}

### 正式强 backbone 对比

{md_table(["方法", "test object RMSE", "OOD object RMSE", "说明"], formal_rows)}

### 我们方法内部消融

{md_table(["分支", "test object RMSE", "OOD object RMSE"], ablation_rows)}

{significance_section}

## 遇到的问题与处理

- 旧实验中直接 depth residual diffusion 不稳定。处理方式是把扩散目标从直接修 depth residual，改为生成/修正 x 相位后验候选，再交给 anchor 与 selector/RCPC 使用。
- 新自建数据目标是 `depth_z`，不同于旧 v1 的 `wall_normal_height`，不能直接复用旧 loader 或直接横向比较旧指标。处理方式是新增独立入口和独立归一化/评估口径。
- 全分辨率 480x640 训练时 batch size 4 对强 backbone 有显存压力。处理方式是 batch size 2 + gradient accumulation 2，保持有效 batch 近似不变。
- GPU 利用率在每 5 epoch validation/checkpoint 时会下降。正式队列采用顺序训练、num_workers 和 feature cache，训练段 GPU 利用率稳定；validation 阶段下降属于正常评估开销。
- 61-64 是不同材质 OOD，普通 test 与 OOD 不能混成一个数。处理方式是在所有表中同时报告普通 test 与 OOD 61-64。
- Word 文档曾出现中文问号。此次 Markdown 使用 UTF-8 保存，Word 使用 python-docx 直接生成 `.docx`，避免终端编码污染。
- 严格相机反投影 PLY 与可视化标准点云存在轻微不齐。严格 PLY 用 calibration 反投影，只用于几何检查；visual z10/视觉对齐版本只看形状，不作为指标。

## 可视化结果

![Formal comparison]({rel(assets['formal'])})

![Selector ablation]({rel(assets['ablation'])})

![Quick screening]({rel(assets['quick'])})

![2D visual comparison]({rel(assets['visual'])})

图中各分支含义：

- Direct base：Attention U-Net 直接预测深度。
- X phase candidate：由单帧预测 x 相位证据转成的深度候选。
- Base+x anchor：Direct base 与 X phase candidate 的平均锚点。
- Phase posterior diffusion candidate：扩散后验产生的 refined x phase/depth 候选。
- Rule final：规则选择/融合后的最终结果，OOD 更稳。
- MLP final：学习式可靠性选择器输出的最终结果，普通 test 最好。
- True x phase oracle：使用真实 x 相位的诊断上界，不是合法 test-time 方法。

## 三维重建对比

严格相机反投影 PLY 文件位于：

`{PAPER_READY / 'pointcloud_obj061_pose02_camera_projective'}`

包括 GT、base+x anchor、Rule final、MLP final 四个版本。它们使用 calibration 中的 3x4 相机矩阵和 0-based 像素反投影，坐标单位沿用数据集 calibration。

![Strict camera PLY preview]({rel(assets['ply_strict'])})

视觉检查版本位于：

`{PAPER_READY / 'pointcloud_obj061_pose02'}`

该版本只用于看形状，不能作为 RMSE 或几何对齐指标。

![Visual PLY preview]({rel(assets['ply_visual'])})

{direct_ply_section}

## 是否足够写论文

本轮结果可以支持论文中的核心实验表，但表述要克制：

- 可以写：在相同单帧输入约束下，phase posterior evidence + anchor + selector/RCPC 相比强 direct backbone 更稳，尤其在异材质 OOD 上更明显。
- 不建议写：扩散直接修 depth residual 就能稳定提升。当前结果不支持这个强 claim。
- 不建议写：绝对 SOTA。当前对比覆盖了常见强 backbone 和 FPP-style multitask proxy，但不是所有结构光最新专用方法。
- 40 epoch 的快速筛选只用于选 baseline；最终主表应使用本报告中的 80 epoch、3 seeds、best-val checkpoint 结果。

## 本地结果文件

- 汇总根目录：`{ROOT}`
- 本报告 Markdown：`{REPORT_MD}`
- 本报告 Word：`{REPORT_DOCX}`
- 指标 JSON：`{SUMMARY_JSON}`
- 指标 CSV：`{SUMMARY_CSV}`
- 正式 direct 结果：`{FORMAL_DIRECT}`
- 正式 Attention U-Net + ours selector：`{FORMAL_SELECTOR}`
- 最佳方法 2D 可视化与 PLY：`{PAPER_READY}`

## 来源代码位置

{md_table(["用途", "本地代码路径"], code_rows)}

服务器上对应代码目录为：

`/root/autodl-tmp/diffusion_fpp_v5`

其中 direct baseline PLY 导出脚本已同步到服务器：

`/root/autodl-tmp/diffusion_fpp_v5/export_single_frame3d_direct_baseline_ply.py`

## 论文建议写法

建议主线写成：

单帧条纹输入先产生直接深度和 x 相位后验候选；扩散模块不再承担直接 depth residual 修正，而是建模更有结构的 x 相位后验证据；最终用 RCPC/selector 在 base、相位候选、扩散候选之间选择或融合。这个叙述与当前实验一致，也能解释为什么单独 refined candidate 不一定最好，但 final result 能稳定超过强 direct backbone。
"""
    return md


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def set_doc_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(10.5)
    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        st = doc.styles[style_name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_table(doc: Document, headers: list[str], rows: list[list[Any]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = str(h)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)


def add_picture_if_exists(doc: Document, path: Path, width: float = 6.4) -> None:
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_docx(direct: dict[str, Any], selector: dict[str, Any], quick: list[dict[str, Any]], assets: dict[str, Path]) -> None:
    doc = Document()
    set_doc_font(doc)
    doc.add_heading("自建数据集单帧相位后验扩散与 RCPC 实验总结", level=0)
    p = doc.add_paragraph("生成时间：2026-06-19。结果根目录：")
    p.add_run(str(ROOT))

    att = direct["attention_unet"]
    unetpp = direct["unetpp"]
    doc.add_heading("1. 结论摘要", level=1)
    for text in [
        f"正式 3-seed 结果中，Attention U-Net direct 的 test RMSE 为 {f4(att['test']['mean'])} ± {f4(att['test']['std'])} mm，UNet++ direct 为 {f4(unetpp['test']['mean'])} ± {f4(unetpp['test']['std'])} mm。",
        f"我们方法接入 Attention U-Net base 后，MLP final 在普通 test 为 {f4(selector['test']['mlp']['mean'])} ± {f4(selector['test']['mlp']['std'])} mm，Rule final 在 61-64 OOD 为 {f4(selector['ood']['rule']['mean'])} ± {f4(selector['ood']['rule']['std'])} mm。",
        "扩散模块更适合作为 x 相位后验候选生成器，而不是直接 depth residual 修正器；最终增益来自 phase evidence、anchor 与 selector/RCPC 的组合。",
        "80 epoch + best-val checkpoint + 3 seeds 足够作为当前论文主表；40 epoch 只作为快速筛选。",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("2. 数据与训练口径", level=1)
    for text in [
        "数据目标为 depth_z；普通 test 共 31 个样本，OOD 61-64 共 12 个异材质样本。",
        "合法 test-time 输入为 input_vertical_0120.bmp 单帧条纹及其单帧派生证据；真实 phase_x/phase_y 只用于监督、诊断或 oracle。",
        "主指标为 object-mask RMSE，单位 mm；valid-mask RMSE 作为辅助指标。",
        "正式 direct backbone 使用 Attention U-Net 和 UNet++，seeds 0/1/2，80 epoch，480x640，batch size 2，accumulation 2，eval every 5 epoch，保存 best-val checkpoint。",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("3. 正式结果", level=1)
    formal_rows = [
        ["Attention U-Net direct", f"{f4(att['test']['mean'])} ± {f4(att['test']['std'])}", f"{f4(att['ood']['mean'])} ± {f4(att['ood']['std'])}"],
        ["UNet++ direct", f"{f4(unetpp['test']['mean'])} ± {f4(unetpp['test']['std'])}", f"{f4(unetpp['ood']['mean'])} ± {f4(unetpp['ood']['std'])}"],
        ["Ours: base+x anchor", f"{f4(selector['test']['anchor']['mean'])} ± {f4(selector['test']['anchor']['std'])}", f"{f4(selector['ood']['anchor']['mean'])} ± {f4(selector['ood']['anchor']['std'])}"],
        ["Ours: Rule final", f"{f4(selector['test']['rule']['mean'])} ± {f4(selector['test']['rule']['std'])}", f"{f4(selector['ood']['rule']['mean'])} ± {f4(selector['ood']['rule']['std'])}"],
        ["Ours: MLP final", f"{f4(selector['test']['mlp']['mean'])} ± {f4(selector['test']['mlp']['std'])}", f"{f4(selector['ood']['mlp']['mean'])} ± {f4(selector['ood']['mlp']['std'])}"],
    ]
    add_table(doc, ["方法", "test object RMSE", "OOD object RMSE"], formal_rows)
    add_picture_if_exists(doc, assets["formal"])
    add_picture_if_exists(doc, assets["ablation"])

    sig_rows_all = significance_rows()
    sig_rows = []
    for split, baseline, candidate in [
        ("test", "attention_unet", "mlp"),
        ("test", "unetpp", "mlp"),
        ("ood", "attention_unet", "rule"),
        ("ood", "unetpp", "rule"),
    ]:
        match = next(
            (
                r
                for r in sig_rows_all
                if r.get("split") == split and r.get("baseline") == baseline and r.get("candidate") == candidate
            ),
            None,
        )
        if match:
            sig_rows.append(
                [
                    split,
                    f"{baseline} -> {candidate}",
                    f"{match['mean_improvement_mm']:.4f}",
                    f"{match['relative_improvement_percent']:.1f}%",
                    f"{match['wins']}/{match['n']}",
                    f"{match['sign_test_p_two_sided']:.4g}",
                ]
            )
    if sig_rows:
        doc.add_heading("4. 样本级显著性分析", level=1)
        doc.add_paragraph("统计方式：先对每个样本的 3 个 seed RMSE 取平均，再做 paired comparison。wins/n 表示候选方法在多少个样本上低于 baseline。")
        add_table(doc, ["split", "paired comparison", "mean improvement mm", "relative", "wins/n", "sign-test p"], sig_rows)
        if "significance" in assets:
            add_picture_if_exists(doc, assets["significance"])

    doc.add_heading("5. 快速筛选结果", level=1)
    quick_rows = [[r["method"], r["epochs"], f4(r["val"]), f4(r["test"]), f4(r["ood"])] for r in quick]
    add_table(doc, ["方法", "epoch", "val RMSE", "test RMSE", "OOD RMSE"], quick_rows)
    add_picture_if_exists(doc, assets["quick"])

    doc.add_heading("6. 可视化与三维重建", level=1)
    doc.add_paragraph("二维可视化包含 direct backbone 与我们方法在同一 OOD 样本上的对比，以及最佳方法 overview。")
    add_picture_if_exists(doc, assets["visual"], width=6.8)
    doc.add_paragraph("严格相机反投影 PLY 使用 calibration 中 3x4 相机矩阵和 0-based 像素反投影，适合检查几何形状；视觉 z10 版本只用于看形状，不作为指标。")
    add_picture_if_exists(doc, assets["ply_strict"], width=6.6)
    add_picture_if_exists(doc, assets["ply_visual"], width=6.6)
    if "ply_direct_strict" in assets:
        doc.add_paragraph("新增 direct baseline 三维点云对比包含 GT、Attention U-Net direct、UNet++ direct、base+x anchor、Rule final 和 MLP final。")
        add_picture_if_exists(doc, assets["ply_direct_strict"], width=6.8)
        add_picture_if_exists(doc, assets["ply_direct_visual"], width=6.8)

    doc.add_heading("7. 遇到的问题与解决", level=1)
    for text in [
        "直接 depth residual diffusion 不稳定，改为 phase posterior evidence，再通过 anchor 与 selector/RCPC 使用。",
        "新数据目标与旧 v1 不同，新增独立 loader、normalization 与评估口径。",
        "强 backbone 全分辨率显存压力较大，使用 batch size 2 + accumulation 2。",
        "validation/checkpoint 阶段 GPU 利用率下降是正常评估开销，训练段利用率稳定。",
        "61-64 OOD 材质差异明显，因此单独报告 OOD 指标。",
        "中文文档编码采用 UTF-8 Markdown 与 python-docx，避免 Word 问号问题。",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("8. 文件位置", level=1)
    for text in [
        f"本地根目录：{ROOT}",
        f"指标 JSON：{SUMMARY_JSON}",
        f"指标 CSV：{SUMMARY_CSV}",
        f"样本级显著性 CSV：{SIGNIFICANCE_CSV}",
        f"样本级均值 CSV：{PER_SAMPLE_MEAN_CSV}",
        f"正式 direct 结果：{FORMAL_DIRECT}",
        f"正式 ours selector：{FORMAL_SELECTOR}",
        f"2D 可视化与 PLY：{PAPER_READY}",
        f"direct baseline PLY：{DIRECT_PLY}",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("9. 来源代码位置", level=1)
    code_rows = [[name, str(path)] for name, path in CODE_PATHS.items()]
    add_table(doc, ["用途", "本地代码路径"], code_rows)
    doc.add_paragraph("服务器上对应代码目录：/root/autodl-tmp/diffusion_fpp_v5")
    doc.add_paragraph(
        "direct baseline PLY 导出脚本已同步到服务器：/root/autodl-tmp/diffusion_fpp_v5/export_single_frame3d_direct_baseline_ply.py"
    )

    doc.save(REPORT_DOCX)


def main() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    set_plot_style()
    direct = aggregate_direct()
    quick = aggregate_quick()
    selector = selector_aggregate()
    write_metric_files(direct, selector, quick)
    assets = {
        "formal": plot_formal_comparison(direct, selector),
        "ablation": plot_selector_ablation(selector),
        "quick": plot_quick_screening(quick),
        "visual": make_visual_contact_sheet(),
    }
    if (ASSET_DIR / "per_sample_paired_improvement.png").exists():
        assets["significance"] = ASSET_DIR / "per_sample_paired_improvement.png"
    assets.update(make_pointcloud_previews())
    md = build_markdown(direct, selector, quick, assets)
    REPORT_MD.write_text(md, encoding="utf-8")
    build_docx(direct, selector, quick, assets)
    print(json.dumps({
        "root": str(ROOT),
        "markdown": str(REPORT_MD),
        "docx": str(REPORT_DOCX),
        "summary_json": str(SUMMARY_JSON),
        "summary_csv": str(SUMMARY_CSV),
        "assets": {k: str(v) for k, v in assets.items()},
    }, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
