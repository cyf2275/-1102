# -*- coding: utf-8 -*-
from __future__ import annotations

import csv
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path.cwd() / "cloud_results" / "A_20260614_single_frame3d_physics_diffusion"
OOD = ROOT / "ood61_64_eval"
OUT = ROOT / "single_frame3d_physics_diffusion_experiment_report.docx"
FALLBACK_OUT = ROOT / "single_frame3d_physics_diffusion_experiment_report_with_ood61_64.docx"
FIG = OOD / "ood61_64_direct_rmse_bar.png"


NAMES = {
    "raw": "仅单帧图像",
    "raw_xy": "单帧图像+坐标",
    "raw_single_phys": "单帧图像+物理派生特征",
    "teacher_aux": "物理特征+训练期相位辅助",
}


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(encoding="utf-8", newline="") as f:
        return list(csv.DictReader(f))


def fmt(x: object) -> str:
    return f"{float(x):.4f}"


def gain(base: float, value: float) -> str:
    return f"{(base - value) / base * 100.0:+.2f}%"


def set_run_font(run, size=10.5, bold=False, color=None, font="宋体") -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_para(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run)
    return p


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text: object, bold=False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    set_run_font(run, size=9, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[object]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        set_cell(table.rows[0].cells[i], header, bold=True)
        shade_cell(table.rows[0].cells[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=9, color=RGBColor(90, 90, 90))


def make_figure(direct_rows: list[dict[str, str]]) -> None:
    order = ["raw", "raw_xy", "raw_single_phys", "teacher_aux"]
    values = [float(next(r for r in direct_rows if r["config"] == cfg)["object_rmse_mean"]) for cfg in order]
    labels = ["single", "single+xy", "single+physics", "physics+aux"]
    colors = ["#7A8799", "#8AA8D8", "#4C9A6A", "#D59A44"]
    plt.figure(figsize=(7.2, 4.0))
    xs = np.arange(len(values))
    bars = plt.bar(xs, values, color=colors, edgecolor="#333333", linewidth=0.7)
    plt.xticks(xs, labels, rotation=12, ha="right")
    plt.ylabel("Object RMSE")
    plt.title("OOD objects 61-64")
    plt.grid(axis="y", alpha=0.25)
    for bar, val in zip(bars, values):
        plt.text(bar.get_x() + bar.get_width() / 2, val + 0.035, f"{val:.3f}", ha="center", va="bottom", fontsize=9)
    plt.tight_layout()
    FIG.parent.mkdir(parents=True, exist_ok=True)
    plt.savefig(FIG, dpi=180)
    plt.close()


def remove_existing_section(doc: Document) -> None:
    # Keep reruns idempotent by removing a previously appended OOD section.
    paragraphs = doc.paragraphs
    start = None
    for i, para in enumerate(paragraphs):
        if para.text.strip().startswith("7. 61-64 异材质外推补充测试"):
            start = i
            break
    if start is None:
        return
    body = doc._body._element
    for element in list(body)[start:]:
        body.remove(element)


def main() -> None:
    direct = read_csv(OOD / "external_eval_direct_aggregated.csv")
    residual = read_csv(OOD / "external_eval_residual_aggregated.csv")
    per_object = read_csv(OOD / "external_eval_direct_per_object.csv")
    make_figure(direct)

    doc = Document(OUT)
    remove_existing_section(doc)

    doc.add_heading("7. 61-64 异材质外推补充测试", level=1)
    add_para(
        doc,
        "用户补充的 obj061 到 obj064 为与前 60 个对象材质差异明显的样本。"
        "本节将它们作为独立外推测试集重新评估已有 checkpoint，共 12 个姿态，"
        "不参与训练，也不改变原 352/80/31 划分。主指标仍为 object mask 区域 RMSE。"
    )

    base = float(next(r for r in direct if r["config"] == "raw")["object_rmse_mean"])
    direct_table = []
    for cfg in ["raw", "raw_xy", "raw_single_phys", "teacher_aux"]:
        row = next(r for r in direct if r["config"] == cfg)
        value = float(row["object_rmse_mean"])
        direct_table.append([
            NAMES[cfg],
            row["seeds"],
            fmt(row["object_rmse_mean"]),
            fmt(row["object_rmse_std"]),
            fmt(row["valid_rmse_mean"]),
            "基准" if cfg == "raw" else gain(base, value),
        ])
    add_table(doc, ["方法", "种子", "object RMSE", "std", "valid RMSE", "相对仅单帧"], direct_table)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(FIG), width=Inches(5.9))
    add_caption(doc, "图 5  61-64 异材质外推测试中不同直接重建方法的 object RMSE。图内英文标签依次对应：仅单帧、单帧+坐标、单帧+物理特征、物理特征+训练期相位辅助。")

    add_para(
        doc,
        "结果显示，61-64 异材质样本上单帧物理派生特征的收益非常明显："
        "相对仅单帧图像，object RMSE 从 2.1337 降至 1.6252，改善约 23.83%。"
        "这与原测试集上约 2.92% 的边缘趋势形成互补证据，说明之前对物理输入的笼统否定不合理。"
        "训练期相位辅助仍未超过纯物理特征，说明辅助监督不是这批样本的主要收益来源。"
    )

    residual_table = []
    for cfg in ["raw", "raw_single_phys", "teacher_aux"]:
        row = next(r for r in residual if r["config"] == cfg)
        residual_table.append([
            NAMES[cfg],
            fmt(row["base_object_rmse_mean"]),
            fmt(row["posterior_mean_object_rmse_mean"]),
            fmt(row["posterior_gate_object_rmse_mean"]),
            f"{float(row['gate_gain_percent']):+.2f}%",
        ])
    add_table(doc, ["基础方法", "基础重建 RMSE", "扩散均值 RMSE", "扩散门控 RMSE", "门控收益"], residual_table)
    add_para(
        doc,
        "扩散后验在异材质外推集上的收益仍然很小：物理特征基础上的门控后验仅改善约 0.11%，"
        "训练期相位辅助基础上约 0.90%，均低于 2% 判定线。"
        "因此当前证据支持物理派生输入，而不支持把扩散 posterior 写成主要增益。"
    )

    per_object_table = []
    for oid in [61, 62, 63, 64]:
        vals = {
            row["config"]: float(row["object_rmse_mean"])
            for row in per_object
            if int(row["object_id"]) == oid
        }
        per_object_table.append([
            oid,
            fmt(vals["raw"]),
            fmt(vals["raw_xy"]),
            fmt(vals["raw_single_phys"]),
            fmt(vals["teacher_aux"]),
        ])
    add_table(doc, ["对象", "仅单帧", "单帧+坐标", "单帧+物理特征", "物理+训练期相位辅助"], per_object_table)
    add_para(
        doc,
        "分对象结果显示，obj061、obj062 和 obj064 中物理特征明显降低误差；"
        "obj063 是异常点，加入物理特征后误差从约 4.28 升至约 6.64。"
        "这提示 obj063 需要单独可视化排查，重点检查材质导致的物理特征失配、标签深度范围和 object mask。"
    )
    add_para(
        doc,
        f"61-64 补充测试的 CSV、JSON 与独立 Markdown 报告保存在：{OOD}"
    )

    try:
        doc.save(OUT)
        print(OUT)
    except PermissionError:
        doc.save(FALLBACK_OUT)
        print(FALLBACK_OUT)


if __name__ == "__main__":
    main()
