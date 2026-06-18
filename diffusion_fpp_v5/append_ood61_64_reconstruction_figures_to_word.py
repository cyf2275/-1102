# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path.cwd() / "cloud_results" / "A_20260614_single_frame3d_physics_diffusion"
DOCX = ROOT / "single_frame3d_physics_diffusion_experiment_report_with_ood61_64.docx"
FALLBACK = ROOT / "single_frame3d_physics_diffusion_experiment_report_with_ood61_64_recon.docx"
VIS = ROOT / "ood61_64_eval" / "visualizations"
DIRECT_FIG = VIS / "ood61_64_reconstruction_direct_contact.png"
DIFF_FIG = VIS / "ood61_64_reconstruction_diffusion_physics_contact.png"


def set_run_font(run, size=10.5, bold=False, color=None, font="宋体") -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run)


def add_picture(doc: Document, path: Path, width: float, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    set_run_font(run, size=9, color=RGBColor(90, 90, 90))


def remove_existing_tail(doc: Document) -> None:
    body = doc._body._element
    start = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith("7.1 61-64 重建可视化"):
            start = i
            break
    if start is None:
        return
    for element in list(body)[start:]:
        body.remove(element)


def main() -> None:
    doc = Document(DOCX)
    remove_existing_tail(doc)
    doc.add_heading("7.1 61-64 重建可视化", level=2)
    add_para(
        doc,
        "为补充分对象数值结果，本节展示 obj061、obj062、obj063、obj064 的代表性重建图。"
        "每个预测图均为 3 个随机种子输出的平均结果。直接重建图用于观察物理特征是否改变形状与局部误差；"
        "扩散图用于观察 residual posterior 在物理特征基础上是否产生有效修正。"
    )
    add_picture(
        doc,
        DIRECT_FIG,
        7.2,
        "图 6  61-64 异材质样本的直接重建可视化。列依次为输入、真值、仅单帧、物理特征、训练期相位辅助、物理特征误差、相位辅助误差。obj063 显示出物理特征分支的明显失配。",
    )
    add_picture(
        doc,
        DIFF_FIG,
        7.2,
        "图 7  61-64 异材质样本上物理特征分支的扩散后验可视化。列依次为输入、真值、物理特征基础重建、扩散均值、扩散门控、基础误差、门控误差、门控相对基础误差变化。",
    )
    add_para(
        doc,
        "从图 6 可以看到，obj061、obj062 和 obj064 中物理特征通常使目标形状更接近真值；"
        "但 obj063 的物理特征和相位辅助分支均出现明显结构性偏差，这解释了分对象表中 obj063 的误差升高。"
        "从图 7 可以看到，扩散门控主要做小幅局部修正，不能系统性修复 obj063 这类基础预测已经失配的样本。"
    )
    try:
        doc.save(DOCX)
        print(DOCX)
    except PermissionError:
        doc.save(FALLBACK)
        print(FALLBACK)


if __name__ == "__main__":
    main()
