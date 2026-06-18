# -*- coding: utf-8 -*-
from pathlib import Path
import csv
import json

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path.cwd() / "cloud_results" / "A_20260614_single_frame3d_physics_diffusion"
VIS = ROOT / "visualizations"
OUT = ROOT / "single_frame3d_physics_diffusion_experiment_report.docx"


def read_csv(path):
    with path.open(encoding="utf-8") as f:
        return list(csv.DictReader(f))


def cfg_name(cfg):
    return {
        "raw": "仅单帧图像",
        "raw_xy": "单帧图像 + 坐标",
        "raw_single_phys": "单帧图像 + 物理派生特征",
        "teacher_aux": "物理特征 + 训练期相位辅助",
    }[cfg]


def fmt(x):
    return f"{float(x):.4f}"


def pct(x):
    return f"{float(x):.2f}%"


def set_run_font(run, size=10.5, bold=False, color=None, font="宋体"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run)
    return p


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    set_run_font(run, size=9, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        set_cell(table.rows[0].cells[i], header, bold=True)
        shade_cell(table.rows[0].cells[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value)
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=9, color=RGBColor(90, 90, 90))


def add_picture(doc, path, width, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def setup_styles(doc):
    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "黑体"
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def main():
    direct_rows = read_csv(ROOT / "direct_aggregated_results.csv")
    residual_rows = read_csv(ROOT / "residual_aggregated_results.csv")
    diff_manifest = json.loads((VIS / "diffusion_teacher_aux_seed1_manifest.json").read_text(encoding="utf-8"))

    doc = Document()
    setup_styles(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("自建真实数据集上的单帧物理输入与扩散后验验证实验报告")
    set_run_font(run, size=18, bold=True, font="黑体")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("实验日期：2026-06-14    数据目标：depth_z    数据集：single_frame_3d_dataset_v1_upload_smalltest")
    set_run_font(run, size=10, color=RGBColor(90, 90, 90))

    doc.add_heading("1. 实验目的与核心结论", level=1)
    add_para(doc, "本实验在新的自建真实采集数据集上重新验证两个问题：第一，单帧图像中加入物理派生特征是否有用；第二，在物理输入有效或接近有效时，扩散残差后验是否还能带来额外收益。")
    add_para(doc, "核心结论是：单帧物理派生特征呈现边缘正向趋势，主要评价区域误差相对仅使用单帧图像下降约 2.92%，非常接近预设 3% 判定线，但严格按规则尚不能判为稳定有效。训练期相位辅助没有带来额外增益。扩散门控后验的平均收益为 0.12% 到 0.53%，远低于 2% 判定线，因此只能作为中性或探索性趋势，不能写成明确有效。")
    add_para(doc, "因此，本次结果不支持继续笼统否定物理输入；更准确的表述是：物理输入值得继续在更大真实数据集上复核，但当前证据尚不足以形成强结论。扩散后验在本数据集上未表现出明确超过基础重建模型的收益。")

    doc.add_heading("2. 数据与实验设置", level=1)
    add_para(doc, "数据集共 463 个样本，划分为 352 个训练样本、80 个验证样本和 31 个测试样本。目标量为 depth_z。由于目标定义与旧自建数据中的 wall_normal_height 不同，本实验结果不能与旧数据或 FPP-ML-Bench 的深度误差直接做数值比较。")
    add_para(doc, "输入约束为单帧合法输入：测试时只允许使用 input_vertical_0120.bmp 及其由单帧图像派生出的物理特征。phase_y、phase_x、bc_y、bc_x 仅用于训练期辅助监督或损失权重，不作为测试时输入。")
    add_table(
        doc,
        ["实验组", "测试时输入", "训练/预测目标", "说明"],
        [
            ["仅单帧图像", "条纹灰度图", "直接回归 depth_z", "基础对照组"],
            ["单帧图像 + 坐标", "条纹灰度图 + x/y 坐标", "直接回归 depth_z", "检验坐标先验是否有帮助"],
            ["单帧图像 + 物理派生特征", "条纹灰度图 + Hilbert/DWT/梯度/FTP 等单帧特征", "直接回归 depth_z", "检验物理输入是否带来收益"],
            ["物理特征 + 训练期相位辅助", "测试输入同物理特征组", "主任务 depth_z，辅助预测相位 sin/cos", "相位只用于训练期监督"],
            ["扩散残差后验", "冻结的基础重建结果 + 对应输入特征", "学习有界残差修正", "比较基础重建、后验均值和门控后验"],
        ],
    )

    doc.add_heading("3. 定量结果", level=1)
    add_para(doc, "主指标为测试集物体区域 RMSE，辅助参考完整有效区域 RMSE。下表中的均值和标准差均基于 3 个随机种子。")
    order = ["raw", "raw_xy", "raw_single_phys", "teacher_aux"]
    direct_table = []
    for cfg in order:
        row = next(item for item in direct_rows if item["config"] == cfg)
        direct_table.append([cfg_name(cfg), row["seeds"], fmt(row["object_rmse_mean"]), fmt(row["object_rmse_std"]), fmt(row["valid_rmse_mean"])])
    add_table(doc, ["直接重建方法", "种子", "物体区域 RMSE 均值", "标准差", "有效区域 RMSE 均值"], direct_table)
    add_picture(doc, VIS / "direct_rmse_bar.png", 5.9, "图 1  直接重建方法的测试集物体区域 RMSE。图内英文标签依次对应：仅单帧图像、单帧+坐标、单帧+物理特征、物理特征+训练期相位辅助。")
    add_para(doc, "单帧物理派生特征相对仅单帧图像的误差下降为 2.9186%，接近但未达到预设 3% 判定线。训练期相位辅助相对物理派生特征下降为 -0.3077%，说明本次没有稳定额外收益。单独加入坐标反而明显变差，提示可用信息更可能来自条纹物理结构，而不是简单空间位置。")

    residual_table = []
    for cfg in ["raw", "raw_single_phys", "teacher_aux"]:
        row = next(item for item in residual_rows if item["config"] == cfg)
        residual_table.append([cfg_name(cfg), row["seeds"], fmt(row["base_object_rmse_mean"]), fmt(row["posterior_mean_object_rmse_mean"]), fmt(row["posterior_gate_object_rmse_mean"]), pct(row["gate_gain_percent"])])
    add_table(doc, ["扩散后验配置", "种子", "基础重建 RMSE", "后验均值 RMSE", "门控后验 RMSE", "门控增益"], residual_table)
    add_picture(doc, VIS / "residual_gain_bar.png", 5.8, "图 2  扩散门控后验相对基础重建的增益。红色虚线为 2% 判定线。")
    add_para(doc, "扩散后验的后验均值在多数组合中反而更差，门控后验能够避免部分错误修正，但平均收益只有 0.12% 到 0.53%。这说明门控策略起到了保守保护作用，但扩散残差本身没有学到足够稳定的改进。")

    doc.add_heading("4. 重建可视化", level=1)
    add_para(doc, "图 3 展示直接重建结果。每一行是一个测试样本，列依次为输入条纹、真值深度、仅单帧重建、物理特征重建、训练期相位辅助重建，以及物理特征组的绝对误差。")
    add_picture(doc, VIS / "reconstruction_direct_contact_seed2.png", 7.1, "图 3  直接重建可视化。物理特征组在部分样本上能改善形状连续性和局部误差，但在困难样本上仍会失败。")
    add_para(doc, "图 4 展示扩散后验可视化。该图选择了扩散门控实际发生变化的一组结果，包含改善明显、轻微改善和变差的样本。列依次为输入、真值、基础重建、扩散后验均值、扩散门控、基础误差、门控误差和误差变化。误差变化中蓝色表示门控后误差降低，红色表示误差升高。")
    add_picture(doc, VIS / "diffusion_teacher_aux_seed1_contact.png", 7.2, "图 4  扩散后验门控可视化。扩散能在部分局部区域降低误差，但也会在部分样本或区域引入错误修正，因此总体收益很小。")

    sample_rows = []
    for item in diff_manifest["selected"]:
        sample_rows.append([
            item["sample_id"],
            f"obj{int(item['object_id']):03d}/pose{int(item['pose_id']):02d}",
            fmt(item["base_rmse"]),
            fmt(item["gate_rmse"]),
            f"{float(item['gain']):+.4f}",
            f"{float(item['accepted_fraction']):.3f}",
        ])
    add_table(doc, ["样本", "对象/姿态", "基础 RMSE", "门控 RMSE", "变化量", "门控接受比例"], sample_rows)
    add_para(doc, "可视化样本表明，扩散门控并非完全无效：它确实能在个别样本上降低 0.08 到 0.15 mm 左右的物体区域 RMSE。但这种改进不稳定，也存在误差上升样本，因此平均后不构成强证据。")

    doc.add_heading("5. 结论边界与后续建议", level=1)
    add_para(doc, "严格按预设规则，物理派生特征没有达到 3% 改善阈值，但结果非常接近阈值，不能继续用“物理输入无效”进行笼统否定。")
    add_para(doc, "训练期相位辅助本轮没有额外收益，不能写作测试时相位输入，也不能写作稳定有效的 teacher 机制。")
    add_para(doc, "扩散后验门控的收益低于 2%，更适合写成 pilot 或 neutral trend，而不是主要贡献。")
    add_para(doc, "后续建议优先扩大真实数据规模、检查困难样本的标注、遮罩和深度范围，并尝试更稳健的物理特征筛选或质量门控；扩散部分应等基础物理输入收益稳定后再继续投入。")

    doc.add_heading("6. 文件位置", level=1)
    add_para(doc, f"本地结果目录：{ROOT}")
    add_para(doc, "主要文件包括：single_frame3d_physics_diffusion_summary.json、direct_aggregated_results.csv、residual_aggregated_results.csv、visualizations 文件夹，以及本 Word 文档。")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
